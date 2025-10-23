import os
import json
import streamlit as st
import numpy as np
import tensorflow as tf
from groq import Groq
from docx import Document
from docx.shared import Inches
import tempfile
from PIL import Image
from google_translate import translations


def Home(lang_code):

    t = translations[lang_code]

    # Header
    st.title ( t["home-page"] )
    st.subheader ( t["subheader"] )
    st.write ( t["features"] )

    st.markdown("---")

    model = tf.keras.models.load_model('best_tomato_model.h5')

    # Classes for diseases
    classes = ['Bacterial_spot', 'Early_blight', 'Late_blight', 'Leaf_Mold', 'No_tomato_leaf', 'Septoria_leaf_spot',
               'Spider_mites Two-spotted_spider_mite', 'Target_Spot', 'Tomato_Yellow_Leaf_Curl_Virus',
               'Tomato_mosaic_virus', 'Healthy', 'powdery_mildew']

    # Function to preprocess image
    def preprocess_image(img) :
        img = img.resize ( (256, 256) )
        img_array = np.array ( img )
        img_array = np.expand_dims ( img_array, axis=0 )
        img_array = img_array / 255.0
        return img_array

    # Tomato Disease Diagnosis Section
    st.title(t["disease_diagnosis"])
    uploaded_file = st.file_uploader(t["upload_prompt"], type=["jpg", "jpeg", "png"])

    if uploaded_file:
        st.image(uploaded_file, caption=t['upload_image'], width=400)

        image = Image.open ( uploaded_file )
        processed_image = preprocess_image ( image )
        prediction = model.predict ( processed_image )

        predicted_class_index = np.argmax ( prediction, axis=1 )[0]
        predicted_class = classes[predicted_class_index]
        confidence = np.max ( prediction ) * 100

        try:
            # Configure Groq API key
            working_dir = os.path.dirname(os.path.abspath(__file__))
            config_data = json.load(open(f"{working_dir}/config.json"))
            GROQ_API_KEY = config_data["GROQ_API_KEY"]
            os.environ["GROQ_API_KEY"] = GROQ_API_KEY

            client = Groq()

            # Query Groq
            messages = [
                {"role": "system", "content": "You are an expert in tomato diseases and treatment solutions."},
                {"role": "user", "content": f"{t['provide_info'].format(predicted_class=predicted_class)}"},

            ]
            response = client.chat.completions.create(
                model="llama-3.1-8b-instant",
                messages=messages
            )
            solution_info = response.choices[0].message.content.strip()

            if predicted_class == "No_tomato_leaf" :
                st.write ( f"### {t['invalid_image']}" )
                st.write ( f" {t['upload_valid_image']}  " )

            elif 50 >= confidence >= 40 :
                st.write (f"### {t['invalid_image']}" )
                st.write (f"{t['upload_valid_image']}  ")

            elif 40 >= confidence >= 0 :
                st.write ( t["low_quality"] )
                st.write ( t ["upload_clear"] )

            else :
                st.write ( f"### {t['disease_detected']} {t[predicted_class]}" )
                st.write ( f"{t['confidence']} {confidence:.2f}%" )
                st.write ( t["solution_info"] )
                st.write(solution_info)

                # Generate Word Report
                if st.button ( t["generate_doc"] ) :
                    doc = Document ()
                    doc.add_heading ( t["title_doc"], level=1 )

                    # Add Disease Name
                    doc.add_paragraph ( f"**{t['disease_detected']}** {predicted_class}" )

                    # Add Image
                    if image :
                        temp_image_path = tempfile.NamedTemporaryFile ( delete=False, suffix=".png" ).name
                        image.save ( temp_image_path )
                        doc.add_picture ( temp_image_path, width=Inches ( 4.5 ) )

                    # Add Solution Info
                    doc.add_heading ( t["solution_info"], level=2 )
                    doc.add_paragraph ( solution_info )

                    # Save Word File
                    temp_doc_path = tempfile.NamedTemporaryFile ( delete=False, suffix=".docx" ).name
                    doc.save ( temp_doc_path )

                    with open ( temp_doc_path, "rb" ) as doc_file :
                        st.download_button ( t["download_report"], doc_file, "Diagnosis_Report.docx" )

        except Exception as e:
            st.error(f"Error: {e}")

    else:
        st.write(t["upload_valid_image"])

    st.markdown ( "---" )
    st.header ( t["offer"] )
    col1, col2 = st.columns ( [2, 1] )
    with col1 :
        st.subheader ( t["buy_seeds"] )
        st.write ( t["seed_benefits"] )
        st.write ( t["why_choose"] )
        st.write ( f"### {t['shop1']}" )

    with col2 :
        ad_image = "./image/tomato_seeds.png"
        st.image ( ad_image, caption=t["buy_seeds"], use_container_width=True )

    st.write("")
    st.write("")

    # Footer
    st.markdown ( f"## {t['footer_title']}" )
    st.write ( t["footer_description"] )
    st.markdown ( f"#### {t['footer_closing']}" )
    st.markdown ( "---" )
