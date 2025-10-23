import streamlit as st
from docx import Document
from docx.shared import Inches
import tempfile
from PIL import Image


# Function to generate a Word file
def generate_word(image, disease_name, solution_info):
    doc = Document()

    # Title
    doc.add_heading("🍅 Tomato Disease Diagnosis Report", level=1)

    # Add Image
    doc.add_paragraph("📷 Uploaded Image:")
    image_path = tempfile.NamedTemporaryFile(delete=False, suffix=".png").name
    image.save(image_path)
    doc.add_picture(image_path, width=Inches(4.5))

    # Add Disease Name
    doc.add_paragraph("\n🦠 Disease Name: ")
    doc.add_heading(disease_name, level=2)

    # Add Solution Info
    doc.add_paragraph("\n💡 Solution & Treatment:")
    doc.add_paragraph(solution_info)

    # Save Word file
    word_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
    doc.save(word_path)

    return word_path

# Streamlit UI
st.title("🍅 Tomato Leaf Disease Detection")

# Upload Image
uploaded_file = st.file_uploader("📂 Upload a Tomato Leaf Image", type=["jpg", "png", "jpeg"])

if uploaded_file:
    image = Image.open(uploaded_file)
    st.image(image, caption="Uploaded Image", use_column_width=True)

    # Simulating Disease Detection (Replace this with your ML model output)
    predicted_disease = "Early Blight"  # Example output
    solution_info = "Use copper-based fungicides and remove infected leaves."

    st.success(f"🦠 **Detected Disease:** {predicted_disease}")

    # Generate Word file
    word_path = generate_word(image, predicted_disease, solution_info)

    # Provide Download Button
    with open(word_path, "rb") as file:
        st.download_button(
            label="📄 Download Diagnosis Report",
            data=file,
            file_name="tomato_disease_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
